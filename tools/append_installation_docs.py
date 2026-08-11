from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "HephaestusWorkbench_Formal_Document_Package_v1.0"


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def append_sds() -> None:
    path = DOCS / "02_SDS_软件设计说明书完整版.docx"
    document = Document(path)
    if any("安装部署与初始化流程" in paragraph.text for paragraph in document.paragraphs):
        return

    document.add_page_break()
    document.add_heading("5. 安装部署与初始化流程", level=1)
    document.add_heading("5.1 安装包与目录边界", level=2)
    document.add_paragraph("安装包采用 .NET 8 self-contained Windows 安装器，默认安装到 C:\\Program Files\\HephaestusWorkbench。程序文件与用户数据严格分离，用户数据默认位于 Documents\\HephaestusWorkbenchData。")
    add_bullets(document, [
        "程序目录保存 HephaestusWorkbench.exe、运行库、WebView2 组件和 PluginSeed。",
        "用户数据目录保存 Database、Cases、Reports、Plugins、Logs、Temp、Config 和 Inbox。",
        "现有 Cases\\<CaseId>\\Report 路径保持不变，不在安装或升级时迁移案例数据。",
    ])

    document.add_heading("5.2 环境检测", level=2)
    document.add_paragraph("安装器检查 Windows 10/11 x64 和 Microsoft Edge WebView2 Runtime。主程序为 self-contained 发布，因此不要求目标机器预装 .NET 8 Desktop Runtime。缺少 WebView2 时，安装器优先执行随包提供的官方 x64 安装程序；无法安装时显示中文错误并停止安装。")

    document.add_heading("5.3 首次运行初始化", level=2)
    document.add_paragraph("首次启动进入五步向导：欢迎、数据目录、日志监控目录、插件目录和初始化完成。初始化过程幂等，任一步骤失败后可以重试，不会覆盖已有案例和报告。")
    add_bullets(document, [
        "创建基础目录并初始化 SQLite 数据库。",
        "默认监控数据目录下的 Inbox；MonitorPaths 支持多个目录。",
        "登记内置日志分析插件。",
        "写入 Config\\appsettings.json、plugins.json 和 workspace.json。",
        "使用 LocalAppData\\HephaestusWorkbench\\bootstrap.json 作为新版本数据根目录指针，不读取或迁移旧版产品数据。",
    ])

    document.add_heading("5.4 正常启动顺序", level=2)
    document.add_paragraph("加载 bootstrap 指针 → 初始化目录和数据库 → 读取当前设置 → 登记内置插件 → 启动多目录日志监控 → 恢复报告会话 → 进入 Dashboard。")

    document.add_heading("5.5 升级、卸载与恢复", level=2)
    add_bullets(document, [
        "升级前将 Database\\workbench.db 备份到 Backups\\upgrade-<timestamp>，暂存解压后替换程序目录。",
        "升级只覆盖程序目录，保留用户数据、案例、报告和插件；版本降级会被阻止。",
        "卸载默认删除程序和快捷方式并保留用户数据，用户明确确认后才删除日志、案例和报告。",
        "数据库损坏时先备份旧文件，再提示创建新数据库；插件清单或入口异常只隔离问题插件，不阻止主程序启动。",
    ])
    document.save(path)


def append_tdd() -> None:
    path = DOCS / "03_TDD_技术详细设计完整版.docx"
    document = Document(path)
    if any("安装器与配置实现" in paragraph.text for paragraph in document.paragraphs):
        return

    document.add_page_break()
    document.add_heading("5. 安装器与配置实现", level=1)
    document.add_heading("5.1 配置模型", level=2)
    document.add_paragraph("WorkbenchConfigurationService 负责三个 JSON 文件的读取、默认值校验和原子写入：")
    add_bullets(document, [
        "appsettings.json：Theme、MaxReportTabs、AutoRestoreReports。",
        "workspace.json：DataPath、MonitorPaths。路径保存为绝对路径并去重。",
        "plugins.json：插件 Id、Version 和 Enabled；插件文件仍由 Plugins 目录发现。",
    ])
    document.add_paragraph("首次生成配置时使用当前数据目录的默认设置；已有 JSON 配置优先保留。写入采用临时文件加替换，避免中断产生半份配置。旧版产品数据不会被自动读取或迁移。")

    document.add_heading("5.2 首次运行向导", level=2)
    document.add_paragraph("FirstRunWizard 使用 WPF MVVM。ViewModel 只收集数据并报告进度，WorkbenchInitializationService 负责创建目录、建库、迁移配置和登记内置插件。LogInboxService 使用多个 FileSystemWatcher 聚合监控目录中的日志文件。")

    document.add_heading("6. 安装、升级和卸载流程", level=1)
    document.add_heading("6.1 安装器", level=2)
    document.add_paragraph("HephaestusWorkbench.Setup 为 Windows x64 self-contained 单文件程序，使用 app.manifest 请求管理员权限。安装器从嵌入式 Payload.zip 解压程序，创建桌面和开始菜单快捷方式，并登记 Windows 应用卸载信息。")
    document.add_heading("6.2 升级保护", level=2)
    document.add_paragraph("升级前检测当前版本和运行中的 HephaestusWorkbench 进程；发现更高版本时拒绝降级。程序文件先解压到临时目录，再替换安装目录，失败时恢复旧目录。数据库备份写入用户数据目录的 Backups 子目录。")
    document.add_heading("6.3 卸载", level=2)
    document.add_paragraph("Windows 卸载入口调用已安装主程序的 --uninstall 模式。用户数据默认保留；确认删除时只允许删除解析后的数据根目录，拒绝删除磁盘根目录，并在当前进程退出后删除程序目录。")

    document.add_heading("7. 测试与验收", level=1)
    add_bullets(document, [
        "配置文件创建、重复初始化和原子写入。",
        "多个日志目录的扫描、聚合、去重和 watcher 切换。",
        "安装包版本读取、暂存替换和路径安全校验。",
        "数据库损坏备份恢复、插件异常隔离和 WebView2 缺失提示。",
        "全新安装、二次启动、升级保留数据以及卸载保留/删除数据。",
    ])
    document.save(path)


if __name__ == "__main__":
    append_sds()
    append_tdd()
